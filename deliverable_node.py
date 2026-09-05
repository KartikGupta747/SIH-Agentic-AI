try:
    from docx import Document
except ImportError:
    import sys
    sys.exit("The 'python-docx' library is required. Please install it using: pip install python-docx")

from graph_state import WorkbenchState
import json

def deliverable_node(state: WorkbenchState) -> dict:
    if "deliverable" not in state.get("active_plan", []):
        return {"final_deliverable_path": "None"}
        
    task_type = state.get("task_type", "CALCULATION")
    
    doc = Document()
    
    if task_type == "CALCULATION":
        doc.add_heading('MRPL Automated Engineering Approval Note', 0)
        
        payload_json = state.get("payload_json", {})
        if not payload_json:
            payload_json = {"Result": "No structured data payload generated."}
            
        doc.add_heading('1. Final Calculated Metrics', level=1)
        for key, value in payload_json.items():
            doc.add_paragraph(f"{key}: {value}", style='List Bullet')
            
        doc.add_heading('2. Execution Audit Trail', level=1)
        audit_log = state.get("audit_log", "No audit log generated.")
        doc.add_paragraph(audit_log)

    elif task_type in ["APPROVAL_VERIFICATION", "POLICY_COMPLIANCE", "PROCUREMENT_VERIFICATION"]:
        doc.add_heading('MRPL Approval Verification Note', 0)
        
        approval_verification = state.get("approval_verification", {})
        if not approval_verification:
            doc.add_paragraph("No approval verification data generated.")
        else:
            doc.add_heading('1. Subject', level=1)
            doc.add_paragraph(approval_verification.get("subject", ""))
            
            doc.add_heading('2. Background / Request', level=1)
            doc.add_paragraph(approval_verification.get("request_summary", ""))
            
            doc.add_heading('3. Financial Details', level=1)
            doc.add_paragraph(approval_verification.get("financial_value", "N/A"))
            
            doc.add_heading('4. Procurement Context', level=1)
            doc.add_paragraph(approval_verification.get("procurement_route", "N/A"))
            
            doc.add_heading('5. Applicable Delegation / Policy', level=1)
            for rule in approval_verification.get("applicable_rules", []):
                doc.add_paragraph(rule, style='List Bullet')
            
            doc.add_heading('6. Evidence Reviewed', level=1)
            for gov_doc in approval_verification.get("governing_documents", []):
                doc.add_paragraph(gov_doc, style='List Bullet')
                
            doc.add_heading('7. Compliance Assessment', level=1)
            doc.add_paragraph(f"Status: {approval_verification.get('compliance_status', 'UNKNOWN')}")
            for finding in approval_verification.get("findings", []):
                doc.add_paragraph(finding, style='List Bullet')
                
            doc.add_heading('8. Exceptions / Gaps', level=1)
            for gap in approval_verification.get("missing_information", []):
                doc.add_paragraph(gap, style='List Bullet')
                
            doc.add_heading('9. Required Approvals', level=1)
            doc.add_paragraph(approval_verification.get("authority_requirement", "N/A"))
            
            doc.add_heading('10. Recommendation', level=1)
            doc.add_paragraph(approval_verification.get("approval_recommendation", ""))
            
    else:
        doc.add_heading('MRPL Analysis Report', 0)
        doc.add_heading('1. User Query', level=1)
        doc.add_paragraph(state.get("user_query", ""))
        doc.add_heading('2. Response', level=1)
        doc.add_paragraph(state.get("final_response", "No structured response generated."))

    file_path = "final_approval_note.docx"
    doc.save(file_path)
    
    return {"final_deliverable_path": file_path}
